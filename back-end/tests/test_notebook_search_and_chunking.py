import io
from types import SimpleNamespace
from uuid import uuid4

import pytest
from docx import Document as DocxDocument

import app.notebooks.tools.chunking as chunking_module
from app.core.config import Settings
from app.notebooks.models import Notebook
from app.notebooks.tools.chunking import (
    ChunkingRequest,
    chunk_document,
    chunk_docx,
    chunk_pdf,
)
from app.notebooks.tools.search import search_notebook_chunks
from app.users.models import User


@pytest.fixture
def settings() -> Settings:
    return Settings(
        database_url="sqlite://",
        jwt_secret_key="test-secret-with-at-least-32-bytes",
        jwt_algorithm="HS256",
        embedding_provider="openai_compatible",
        embedding_api_key="key",
    )


def test_docx_chunking_includes_table_text() -> None:
    doc = DocxDocument()
    doc.add_paragraph("Intro paragraph")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Header A"
    table.rows[0].cells[1].text = "Header B"

    buffer = io.BytesIO()
    doc.save(buffer)

    chunks = chunk_docx(
        ChunkingRequest(
            content=buffer.getvalue(),
            filename="sample.docx",
            source="docx-source",
            document_id="doc-1",
        ),
        chunk_size=1000,
        chunk_overlap=200,
    )

    combined = "\n".join(chunk.page_content for chunk in chunks)
    assert "Intro paragraph" in combined
    assert "Header A | Header B" in combined


def test_pdf_chunking_extracts_simple_text(monkeypatch: pytest.MonkeyPatch) -> None:
    class FakePage:
        def extract_text(self) -> str:
            return "PDF text"

    class FakePdf:
        def __init__(self) -> None:
            self.pages = [FakePage()]

        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, tb):
            return False

    monkeypatch.setattr(
        "app.notebooks.tools.chunking.pdfplumber.open", lambda _f: FakePdf()
    )

    chunks = chunk_pdf(
        ChunkingRequest(
            content=b"%PDF-1.7",
            filename="sample.pdf",
            source="pdf-source",
            document_id="doc-2",
        ),
        chunk_size=1000,
        chunk_overlap=200,
    )

    assert any("PDF text" in chunk.page_content for chunk in chunks)


def test_chunk_document_uses_settings_for_docx(monkeypatch: pytest.MonkeyPatch) -> None:
    captured: dict[str, int] = {}

    def fake_chunk_docx(
        request: ChunkingRequest,
        *,
        settings: Settings | None = None,
        chunk_size: int,
        chunk_overlap: int,
    ):
        captured["chunk_size"] = chunk_size
        captured["chunk_overlap"] = chunk_overlap
        return []

    monkeypatch.setitem(chunking_module.ENGINE_BY_EXTENSION, ".docx", fake_chunk_docx)

    request = ChunkingRequest(
        content=b"docx-bytes",
        filename="sample.docx",
        source="docx-source",
        document_id="doc-3",
    )
    settings = Settings(
        database_url="sqlite://",
        jwt_secret_key="test-secret-with-at-least-32-bytes",
        jwt_algorithm="HS256",
        notebook_chunk_size=2048,
        notebook_chunk_overlap=128,
    )

    assert chunk_document(request, settings) == []
    assert captured == {"chunk_size": 2048, "chunk_overlap": 128}


def test_semantic_chunking_success(
    monkeypatch: pytest.MonkeyPatch, settings: Settings
) -> None:
    text = "First sentence about cats. Second sentence about dogs. Third sentence about math."
    # 3 sentences -> 3 embeddings
    mock_vectors = [
        [1.0] * 1536,
        [1.0] * 1536,
        [0.0] * 1536,
    ]

    def fake_embed_texts(texts: list[str], s: Settings) -> list[list[float]]:
        return [mock_vectors[i % len(mock_vectors)] for i in range(len(texts))]

    monkeypatch.setattr(
        "app.notebooks.tools.embeddings.embed_texts",
        fake_embed_texts,
    )

    request = ChunkingRequest(
        content=text.encode("utf-8"),
        filename="sample.txt",
        source="txt-source",
        document_id="doc-5",
    )

    chunks = chunk_document(request, settings)
    assert len(chunks) > 0
    assert any("First sentence" in chunk.page_content for chunk in chunks)


def test_chunk_document_rejects_unsupported_extension() -> None:
    request = ChunkingRequest(
        content=b"plain text",
        filename="sample.csv",
        source="csv-source",
        document_id="doc-4",
    )

    with pytest.raises(ValueError, match=r"Unsupported file type: \.csv"):
        chunk_document(request)


def test_search_postgres_uses_vector_operator(
    monkeypatch: pytest.MonkeyPatch, settings: Settings
) -> None:
    user_id = uuid4()
    doc_id = uuid4()

    class FakeResult:
        def all(self):
            return [(doc_id, "source.txt", 0, "content", {})]

    class FakeSession:
        def __init__(self):
            self.bind = SimpleNamespace(dialect=SimpleNamespace(name="postgresql"))
            self.executed_stmt = None
            self.executed_params = None

        def execute(self, stmt, params):
            self.executed_stmt = str(stmt)
            self.executed_params = params
            return FakeResult()

    fake_session = FakeSession()
    monkeypatch.setattr(
        "app.notebooks.tools.search.embed_texts",
        lambda texts, _settings: [[0.1] * 1536],
    )

    notebook = Notebook(user_id=user_id, name="Notebook", description="", tags=[])
    current_user = User(
        id=user_id, email="user@example.com", hashed_password="hashed-password"
    )

    results = search_notebook_chunks(
        fake_session,
        notebook=notebook,
        current_user=current_user,
        query="q",
        settings=settings,
        top_k=4,
    )

    assert len(results) == 1
    assert fake_session.executed_stmt is not None
    assert "<=>" in fake_session.executed_stmt
    assert fake_session.executed_params is not None
    assert fake_session.executed_params["top_k"] == 4
