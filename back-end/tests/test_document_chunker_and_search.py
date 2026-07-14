import io
from typing import Any
from unittest.mock import AsyncMock
from uuid import uuid4

import pytest
from bson import Binary
from docx import Document as DocxDocument
from langchain_core.documents import Document

import app.notebooks.rag.document_chunker as chunking_module
import app.notebooks.rag.search_service as search_service_module
from app.core.config import Settings
from app.notebooks.models import Notebook, NotebookDocumentChunk
from app.notebooks.rag.document_chunker import (
    ChunkingRequest,
    chunk_document,
)
from app.notebooks.rag.ingestion_service import _order_pdf_chunks
from app.notebooks.rag.query_rewrite_agent import (
    rewrite_query_text,
)
from app.notebooks.rag.search_service import search_notebook_chunks
from app.users.models import User


@pytest.fixture
def settings() -> Settings:
    return Settings(
        database_url="mongodb://localhost:27017/test",
        jwt_secret_key="test-secret-key",
        jwt_algorithm="HS256",
        access_token_expire_minutes=30,
        refresh_token_expire_days=30,
        otp_expire_minutes=10,
        otp_max_attempts=5,
        log_level="DEBUG",
        resend_api_key="",
        cookie_secure=False,
        notebook_chunk_size=100,
        notebook_chunk_overlap=20,
        embedding_dimension=1536,
        embedding_model="text-embedding-3-small",
    )


def test_docx_chunking_includes_table_text(settings: Settings) -> None:
    doc = DocxDocument()
    doc.add_paragraph("Intro paragraph")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Header A"
    table.rows[0].cells[1].text = "Header B"

    buffer = io.BytesIO()
    doc.save(buffer)

    chunks = chunk_document(
        ChunkingRequest(
            content=buffer.getvalue(),
            filename="sample.docx",
            source="docx-source",
            document_id="doc-1",
        ),
        settings,
    )

    combined = "\n".join(chunk.page_content for chunk in chunks)
    assert "Intro paragraph" in combined
    assert "Header A | Header B" in combined


def test_pdf_chunking_extracts_simple_text(
    monkeypatch: pytest.MonkeyPatch, settings: Settings
) -> None:
    monkeypatch.setattr(chunking_module.pymupdf, "open", lambda **kwargs: "fake_doc")
    monkeypatch.setattr(
        chunking_module.pymupdf4llm,
        "to_markdown",
        lambda doc, **kwargs: [{"text": "PDF text"}],
    )

    chunks = chunk_document(
        ChunkingRequest(
            content=b"%PDF-1.7",
            filename="sample.pdf",
            source="pdf-source",
            document_id="doc-2",
        ),
        settings,
    )

    assert any("PDF text" in chunk.page_content for chunk in chunks)


def test_pdf_chunking_strips_image_placeholders(
    monkeypatch: pytest.MonkeyPatch, settings: Settings
) -> None:
    # pymupdf4llm leaves a placeholder wherever it drops an image; we describe
    # images via the vision pipeline instead, so these must not leak into the
    # embedded text (or the model reads them back as "picture omitted").
    markdown = (
        "Intro paragraph.\n\n"
        "**==> picture [246 x 81] intentionally omitted <==**\n\n"
        "Body paragraph.\n\n"
        "==> picture [100.5 x 42.0] <==\n\n"
        "Conclusion."
    )
    monkeypatch.setattr(chunking_module.pymupdf, "open", lambda **kwargs: "fake_doc")
    monkeypatch.setattr(
        chunking_module.pymupdf4llm,
        "to_markdown",
        lambda doc, **kwargs: [{"text": markdown}],
    )

    chunks = chunk_document(
        ChunkingRequest(
            content=b"%PDF-1.7",
            filename="sample.pdf",
            source="pdf-source",
            document_id="doc-img",
        ),
        settings,
    )

    combined = "\n".join(chunk.page_content for chunk in chunks)
    assert "picture" not in combined
    assert "omitted" not in combined
    assert "Intro paragraph." in combined
    assert "Body paragraph." in combined
    assert "Conclusion." in combined


def test_pdf_chunking_assigns_page_numbers(
    monkeypatch: pytest.MonkeyPatch, settings: Settings
) -> None:
    # Long enough per-page text that each page lands in its own chunk given
    # the fixture's chunk_size=100 (short pages would merge into one chunk).
    monkeypatch.setattr(chunking_module.pymupdf, "open", lambda **kwargs: "fake_doc")
    monkeypatch.setattr(
        chunking_module.pymupdf4llm,
        "to_markdown",
        lambda doc, **kwargs: [
            {"text": "Page one content. " * 10},
            {"text": "Page two content. " * 10},
            {"text": "Page three content. " * 10},
        ],
    )

    chunks = chunk_document(
        ChunkingRequest(
            content=b"%PDF-1.7",
            filename="sample.pdf",
            source="pdf-source",
            document_id="doc-pages",
        ),
        settings,
    )

    page_numbers_by_content_prefix = {
        chunk.page_content[:8]: chunk.metadata["page_number"] for chunk in chunks
    }
    assert page_numbers_by_content_prefix["Page one"] == 1
    assert page_numbers_by_content_prefix["Page two"] == 2
    assert page_numbers_by_content_prefix["Page thr"] == 3


def test_order_pdf_chunks_interleaves_images_by_page() -> None:
    text_docs = [
        Document(page_content="page 1 text", metadata={"page_number": 1}),
        Document(page_content="page 2 text", metadata={"page_number": 2}),
    ]
    image_docs = [
        Document(
            page_content="figure on page 1",
            metadata={"chunk_type": "image", "page_number": 1},
        ),
    ]

    ordered = _order_pdf_chunks(text_docs, image_docs)

    assert [doc.page_content for doc in ordered] == [
        "page 1 text",
        "figure on page 1",
        "page 2 text",
    ]


def test_order_pdf_chunks_sorts_missing_page_number_last() -> None:
    text_docs = [Document(page_content="has page", metadata={"page_number": 1})]
    image_docs = [
        Document(page_content="no page", metadata={"chunk_type": "image"}),
    ]

    ordered = _order_pdf_chunks(text_docs, image_docs)

    assert [doc.page_content for doc in ordered] == ["has page", "no page"]


def test_chunk_document_uses_settings_for_docx(monkeypatch: pytest.MonkeyPatch) -> None:
    captured: dict[str, int] = {}

    def fake_split_text(
        text: str,
        *,
        source: str,
        document_id: str,
        chunk_size: int,
        chunk_overlap: int,
        page_offsets: list[int] | None = None,
    ) -> list:
        captured["chunk_size"] = chunk_size
        captured["chunk_overlap"] = chunk_overlap
        return []

    monkeypatch.setitem(
        chunking_module.EXTRACTOR_BY_EXTENSION, ".docx", lambda request: "docx text"
    )
    monkeypatch.setattr(chunking_module, "split_text", fake_split_text)

    request = ChunkingRequest(
        content=b"docx-bytes",
        filename="sample.docx",
        source="docx-source",
        document_id="doc-3",
    )
    settings = Settings(
        database_url="mongodb://localhost:27017/test",
        jwt_secret_key="test-secret-key",
        jwt_algorithm="HS256",
        access_token_expire_minutes=30,
        refresh_token_expire_days=30,
        notebook_chunk_size=2048,
        notebook_chunk_overlap=128,
        embedding_dimension=1536,
        embedding_model="text-embedding-3-small",
    )

    assert chunk_document(request, settings) == []
    assert captured == {"chunk_size": 2048, "chunk_overlap": 128}


def test_chunk_document_rejects_unsupported_extension(settings: Settings) -> None:
    request = ChunkingRequest(
        content=b"plain text",
        filename="sample.csv",
        source="csv-source",
        document_id="doc-4",
    )

    with pytest.raises(ValueError, match=r"Unsupported file type: \.csv"):
        chunk_document(request, settings)


@pytest.mark.anyio
async def test_rewrite_query_text_disabled(
    monkeypatch: pytest.MonkeyPatch, settings: Settings
) -> None:
    monkeypatch.setattr(
        "app.notebooks.rag.query_rewrite_agent.chat_provider_is_configured",
        lambda: True,
    )
    settings.enable_query_rewrite = False

    res = await rewrite_query_text("can you find apples?", settings)
    assert res == "can you find apples?"


@pytest.mark.anyio
async def test_rewrite_query_text_success(
    monkeypatch: pytest.MonkeyPatch, settings: Settings
) -> None:
    from pydantic_ai.models.test import TestModel

    monkeypatch.setattr(
        "app.notebooks.rag.query_rewrite_agent.chat_provider_is_configured",
        lambda: True,
    )

    monkeypatch.setattr(
        "app.notebooks.rag.query_rewrite_agent.resolve_chat_provider",
        lambda: TestModel(custom_output_text="apples"),
    )

    settings.enable_query_rewrite = True
    res = await rewrite_query_text("can you find apples?", settings)
    assert res == "apples"


@pytest.mark.anyio
async def test_rewrite_query_text_failure_fallback(
    monkeypatch: pytest.MonkeyPatch, settings: Settings
) -> None:
    monkeypatch.setattr(
        "app.notebooks.rag.query_rewrite_agent.chat_provider_is_configured",
        lambda: True,
    )

    def failing_resolve():
        raise RuntimeError("LLM error")

    monkeypatch.setattr(
        "app.notebooks.rag.query_rewrite_agent.resolve_chat_provider", failing_resolve
    )

    settings.enable_query_rewrite = True
    res = await rewrite_query_text("can you find apples?", settings)
    assert res == "can you find apples?"


@pytest.mark.anyio
async def test_search_notebook_chunks_filters_by_document_id(
    monkeypatch: pytest.MonkeyPatch, settings: Settings
) -> None:
    captured_pipeline: list[dict[str, Any]] = []

    def fake_aggregate(pipeline: list[dict[str, Any]]) -> AsyncMock:
        captured_pipeline.extend(pipeline)
        cursor = AsyncMock()
        cursor.to_list.return_value = []
        return cursor

    monkeypatch.setattr(NotebookDocumentChunk, "aggregate", fake_aggregate)
    monkeypatch.setattr(
        search_service_module, "embed_text", AsyncMock(return_value=[0.0])
    )
    monkeypatch.setattr(
        search_service_module, "rewrite_query_text", AsyncMock(return_value="query")
    )

    notebook = Notebook(id=uuid4(), user_id=uuid4(), name="Test Notebook")
    user = User(id=uuid4(), email="user@example.com", hashed_password="hashed")
    document_ids = [uuid4(), uuid4()]

    await search_notebook_chunks(
        notebook=notebook,
        current_user=user,
        query="query",
        settings=settings,
        document_ids=document_ids,
    )

    search_filter = captured_pipeline[0]["$vectorSearch"]["filter"]
    assert search_filter["document_id"] == {
        "$in": [Binary.from_uuid(doc_id) for doc_id in document_ids]
    }
    assert search_filter["notebook_id"] == Binary.from_uuid(notebook.id)


@pytest.mark.anyio
async def test_search_notebook_chunks_omits_document_id_filter_by_default(
    monkeypatch: pytest.MonkeyPatch, settings: Settings
) -> None:
    captured_pipeline: list[dict[str, Any]] = []

    def fake_aggregate(pipeline: list[dict[str, Any]]) -> AsyncMock:
        captured_pipeline.extend(pipeline)
        cursor = AsyncMock()
        cursor.to_list.return_value = []
        return cursor

    monkeypatch.setattr(NotebookDocumentChunk, "aggregate", fake_aggregate)
    monkeypatch.setattr(
        search_service_module, "embed_text", AsyncMock(return_value=[0.0])
    )
    monkeypatch.setattr(
        search_service_module, "rewrite_query_text", AsyncMock(return_value="query")
    )

    notebook = Notebook(id=uuid4(), user_id=uuid4(), name="Test Notebook")
    user = User(id=uuid4(), email="user@example.com", hashed_password="hashed")

    await search_notebook_chunks(
        notebook=notebook, current_user=user, query="query", settings=settings
    )

    search_filter = captured_pipeline[0]["$vectorSearch"]["filter"]
    assert "document_id" not in search_filter
