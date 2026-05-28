from __future__ import annotations

import io
import logging
from datetime import UTC, datetime
from pathlib import Path
from uuid import UUID

from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
from sqlalchemy.exc import SQLAlchemyError
from sqlmodel import Session, delete, select

from app.core.config import Settings
from app.notebooks.models import Notebook, NotebookDocument, NotebookDocumentChunk
from app.notebooks.tools.search import embed_texts
from app.users.models import User

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200


def _is_created_event(event_name: str | None) -> bool:
    if not event_name:
        return False
    lowered = event_name.lower()
    return "objectcreated" in lowered or "object_created" in lowered


def register_pending_notebook_document(
    session: Session,
    *,
    notebook: Notebook,
    current_user: User,
    bucket: str,
    key: str,
    filename: str,
    content_type: str | None,
) -> NotebookDocument:
    document = NotebookDocument(
        notebook_id=notebook.id,
        user_id=current_user.id,
        s3_bucket=bucket,
        s3_key=key,
        filename=filename,
        content_type=content_type,
        status="pending",
    )
    session.add(document)
    session.commit()
    session.refresh(document)
    return document


def mark_document_uploaded_and_get_id(
    session: Session,
    *,
    bucket: str | None,
    key: str | None,
    size: int | None,
    event_name: str | None,
) -> UUID | None:
    if not key or not _is_created_event(event_name):
        return None
    statement = select(NotebookDocument).where(NotebookDocument.s3_key == key)
    if bucket:
        statement = statement.where(NotebookDocument.s3_bucket == bucket)
    try:
        document = session.exec(statement).first()
    except SQLAlchemyError:
        logger.warning("Skipping notebook document callback update; table may not exist yet")
        return None
    if document is None:
        return None
    now = datetime.now(UTC)
    document.status = "uploaded"
    document.size = size
    document.error_message = None
    document.updated_at = now
    session.add(document)
    session.commit()
    return document.id


def mark_document_upload_failed(
    session: Session,
    *,
    key: str,
    user_id: UUID,
    error_message: str | None = None,
) -> bool:
    statement = select(NotebookDocument).where(
        NotebookDocument.s3_key == key,
        NotebookDocument.user_id == user_id,
    )
    document = session.exec(statement).first()
    if document is None:
        return False
    if document.status in {"indexed", "processing", "uploaded"}:
        return True
    document.status = "failed"
    document.error_message = (error_message or "Upload failed before object storage accepted the file.")[:4000]
    document.updated_at = datetime.now(UTC)
    session.add(document)
    session.commit()
    return True


def _extract_text_from_bytes(content: bytes, filename: str, content_type: str | None) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {suffix or 'unknown'}")

    if suffix in {".txt", ".md"}:
        return content.decode("utf-8", errors="ignore")

    if suffix == ".pdf":
        reader = PdfReader(io.BytesIO(content))
        return "\n".join((page.extract_text() or "") for page in reader.pages)

    if suffix == ".docx":
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(content))
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)

    raise ValueError(f"Unsupported file type: {suffix}, content_type={content_type}")


def ingest_document_by_id(session: Session, document_id: UUID, settings: Settings) -> None:
    document = session.get(NotebookDocument, document_id)
    if document is None:
        return

    document.status = "processing"
    document.error_message = None
    document.updated_at = datetime.now(UTC)
    session.add(document)
    session.commit()

    try:
        from app.file.service import get_s3_client

        s3_client = get_s3_client(settings)
        obj = s3_client.get_object(Bucket=document.s3_bucket, Key=document.s3_key)
        body = obj["Body"].read()
        extracted_text = _extract_text_from_bytes(body, document.filename, document.content_type)
        if not extracted_text.strip():
            raise ValueError("No extractable text content in document")

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            add_start_index=True,
        )
        split_docs = splitter.create_documents(
            [extracted_text],
            metadatas=[{"source": document.s3_key, "document_id": str(document.id)}],
        )
        chunk_texts = [doc.page_content for doc in split_docs]
        if not chunk_texts:
            raise ValueError("No chunks produced from extracted text")

        embeddings = embed_texts(chunk_texts, settings)

        session.exec(delete(NotebookDocumentChunk).where(NotebookDocumentChunk.document_id == document.id))
        now = datetime.now(UTC)
        for idx, split_doc in enumerate(split_docs):
            session.add(
                NotebookDocumentChunk(
                    document_id=document.id,
                    notebook_id=document.notebook_id,
                    user_id=document.user_id,
                    chunk_index=idx,
                    content=split_doc.page_content,
                    chunk_metadata=split_doc.metadata,
                    embedding=embeddings[idx],
                    created_at=now,
                    updated_at=now,
                )
            )

        document.status = "indexed"
        document.error_message = None
        document.updated_at = now
        session.add(document)
        session.commit()
    except Exception as exc:  # pragma: no cover - error path exercised by tests via status check
        session.rollback()
        logger.exception("Notebook document ingestion failed for %s", document_id)
        document = session.get(NotebookDocument, document_id)
        if document is not None:
            document.status = "failed"
            document.error_message = str(exc)[:4000]
            document.updated_at = datetime.now(UTC)
            session.add(document)
            session.commit()
