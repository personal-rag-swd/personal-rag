from __future__ import annotations

import io
from dataclasses import dataclass
from pathlib import Path
from typing import TYPE_CHECKING

import pdfplumber
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

if TYPE_CHECKING:
    from app.core.config import Settings

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


@dataclass(frozen=True)
class ChunkingRequest:
    content: bytes
    filename: str
    source: str
    document_id: str


class LangChainEmbeddingAdapter(Embeddings):
    def __init__(self, settings: Settings) -> None:
        self.settings = settings

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        if not texts:
            return []
        from app.notebooks.tools.embeddings import embed_texts

        return embed_texts(texts, self.settings)

    def embed_query(self, text: str) -> list[float]:
        from app.notebooks.tools.embeddings import embed_texts

        return embed_texts([text], self.settings)[0]


def split_text(
    text: str,
    *,
    source: str,
    document_id: str,
    settings: Settings | None = None,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
) -> list[Document]:
    if settings is None:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            add_start_index=True,
        )
        return splitter.create_documents(
            [text],
            metadatas=[{"source": source, "document_id": document_id}],
        )

    from langchain_experimental.text_splitter import SemanticChunker

    embeddings = LangChainEmbeddingAdapter(settings)
    splitter = SemanticChunker(
        embeddings=embeddings,
        add_start_index=True,
        breakpoint_threshold_type="percentile",
    )
    return splitter.create_documents(
        [text],
        metadatas=[{"source": source, "document_id": document_id}],
    )


def chunk_pdf(
    request: ChunkingRequest,
    *,
    settings: Settings | None = None,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
) -> list[Document]:
    with pdfplumber.open(io.BytesIO(request.content)) as pdf:
        extracted_text = "\n".join((page.extract_text() or "") for page in pdf.pages)
    return split_text(
        extracted_text,
        source=request.source,
        document_id=request.document_id,
        settings=settings,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )


def chunk_docx(
    request: ChunkingRequest,
    *,
    settings: Settings | None = None,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
) -> list[Document]:
    doc = DocxDocument(io.BytesIO(request.content))
    paragraphs = [
        paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()
    ]
    table_rows: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_rows.append(" | ".join(cells))
    extracted_parts = paragraphs + table_rows
    extracted_text = "\n".join(extracted_parts)
    return split_text(
        extracted_text,
        source=request.source,
        document_id=request.document_id,
        settings=settings,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )


def chunk_text(
    request: ChunkingRequest,
    *,
    settings: Settings | None = None,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
) -> list[Document]:
    extracted_text = request.content.decode("utf-8", errors="ignore")
    return split_text(
        extracted_text,
        source=request.source,
        document_id=request.document_id,
        settings=settings,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )


ENGINE_BY_EXTENSION = {
    ".pdf": chunk_pdf,
    ".docx": chunk_docx,
    ".txt": chunk_text,
    ".md": chunk_text,
}


def chunk_document(
    request: ChunkingRequest, settings: Settings | None = None
) -> list[Document]:
    suffix = Path(request.filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {suffix or 'unknown'}")

    engine = ENGINE_BY_EXTENSION[suffix]
    chunk_size = settings.notebook_chunk_size if settings is not None else 1000
    chunk_overlap = settings.notebook_chunk_overlap if settings is not None else 200
    return engine(
        request,
        settings=settings,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )
