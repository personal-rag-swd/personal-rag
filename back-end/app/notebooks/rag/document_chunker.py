from __future__ import annotations

import bisect
import io
import logging
import re
from dataclasses import dataclass
from pathlib import Path
from typing import TYPE_CHECKING

import pymupdf
import pymupdf4llm
from docx import Document as DocxDocument
from docx.table import Table as DocxTable
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

if TYPE_CHECKING:
    from app.core.config import Settings

logger = logging.getLogger(__name__)

# Extensions handled by chunk_document() (text-extraction engines below).
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
# Canonical IANA image MIME types keyed by bare extension (no leading dot).
# "image/jpg" is NOT valid — vision providers reject it — so "jpg" maps to
# "image/jpeg". This is the single source of truth for both the upload gate
# (IMAGE_EXTENSIONS) and ingestion_service's media-type lookups.
IMAGE_MEDIA_TYPES = {
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "png": "image/png",
    "webp": "image/webp",
    "gif": "image/gif",
}
# Image extensions are accepted for upload but routed to the async vision
# pipeline in ingestion_service, not to chunk_document().
IMAGE_EXTENSIONS = {f".{ext}" for ext in IMAGE_MEDIA_TYPES}


@dataclass(frozen=True)
class ChunkingRequest:
    content: bytes
    filename: str
    source: str
    document_id: str


def split_text(
    text: str,
    *,
    source: str,
    document_id: str,
    chunk_size: int,
    chunk_overlap: int,
    page_offsets: list[int] | None = None,
) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        add_start_index=True,
    )
    docs = splitter.create_documents(
        [text],
        metadatas=[{"source": source, "document_id": document_id}],
    )
    if page_offsets:
        for doc in docs:
            # bisect_right gives the count of page offsets at or before
            # start_index, which is exactly the 1-based page number.
            doc.metadata["page_number"] = bisect.bisect_right(
                page_offsets, doc.metadata["start_index"]
            )
    logger.info("Split %s into %d chunk(s)", source, len(docs))
    return docs


# pymupdf4llm drops every image from its markdown output (we describe images
# separately via the vision pipeline) and leaves a placeholder in the text,
# e.g. "**==> picture [246 x 81] intentionally omitted <==**". Left in, it gets
# embedded as a text chunk and read back to the user as "the picture was
# omitted", so strip it out. Covers both emitted variants (with and without the
# "intentionally omitted" wording), with or without the surrounding bold markers.
_IMAGE_PLACEHOLDER_RE = re.compile(
    r"[ \t]*\*{0,2}==>\s*picture\s*\[[\d.]+\s*x\s*[\d.]+\]"
    r"(?:\s*intentionally omitted)?\s*<==\*{0,2}[ \t]*\n?",
    re.IGNORECASE,
)
# Collapse the runs of blank lines left behind after removing the placeholders.
_EXCESS_BLANK_LINES_RE = re.compile(r"\n{3,}")


def _extract_pdf_text(request: ChunkingRequest) -> tuple[str, list[int]]:
    """Extract PDF text and the character offset each page starts at.

    Uses ``page_chunks=True`` so pages are joined ourselves (rather than
    pymupdf4llm's single blob) — the offsets let ``split_text`` map each
    chunk's ``start_index`` back to the page it came from, so text chunks
    carry ``page_number`` the same way image chunks do.
    """
    doc = pymupdf.open(stream=request.content, filetype="pdf")
    pages = pymupdf4llm.to_markdown(doc, page_chunks=True)
    if not isinstance(pages, list):
        raise TypeError(
            f"pymupdf4llm.to_markdown(page_chunks=True) returned unexpected "
            f"type: {type(pages)}"
        )

    page_texts: list[str] = []
    for page in pages:
        text = page["text"] if isinstance(page, dict) else str(page)
        text = _IMAGE_PLACEHOLDER_RE.sub("", text)
        page_texts.append(_EXCESS_BLANK_LINES_RE.sub("\n\n", text))

    page_offsets: list[int] = []
    cursor = 0
    for text in page_texts:
        page_offsets.append(cursor)
        cursor += len(text) + 2  # +2 accounts for the "\n\n" join separator
    return "\n\n".join(page_texts), page_offsets


def _docx_table_rows(table: DocxTable) -> list[str]:
    rows: list[str] = []
    for row in table.rows:
        # Deduplicate adjacent merged cells that repeat the same text
        seen: set[str] = set()
        cells: list[str] = []
        for cell in row.cells:
            text = cell.text.strip()
            if text and text not in seen:
                seen.add(text)
                cells.append(text)
        if cells:
            rows.append(" | ".join(cells))
    return rows


def _extract_docx_text(request: ChunkingRequest) -> str:
    doc = DocxDocument(io.BytesIO(request.content))
    parts: list[str] = []
    # iter_inner_content yields paragraphs and tables in document order,
    # unlike doc.paragraphs + doc.tables which regroups them by kind.
    for item in doc.iter_inner_content():
        if isinstance(item, DocxTable):
            parts.extend(_docx_table_rows(item))
        elif item.text.strip():
            parts.append(item.text.strip())
    return "\n".join(parts)


def _extract_plain_text(request: ChunkingRequest) -> str:
    return request.content.decode("utf-8", errors="ignore")


# PDFs are dispatched separately (chunk_document below) since _extract_pdf_text
# also returns page offsets; the others are plain str extractors.
EXTRACTOR_BY_EXTENSION = {
    ".docx": _extract_docx_text,
    ".txt": _extract_plain_text,
    ".md": _extract_plain_text,
}


def chunk_document(request: ChunkingRequest, settings: Settings) -> list[Document]:
    suffix = Path(request.filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {suffix or 'unknown'}")

    logger.info(
        "Extracting text: filename=%s, source=%s, size=%d bytes, format=%s, chunk_size=%d",
        request.filename,
        request.source,
        len(request.content),
        suffix,
        settings.notebook_chunk_size,
    )
    page_offsets: list[int] | None = None
    if suffix == ".pdf":
        extracted_text, page_offsets = _extract_pdf_text(request)
    else:
        extracted_text = EXTRACTOR_BY_EXTENSION[suffix](request)
    return split_text(
        extracted_text,
        source=request.source,
        document_id=request.document_id,
        chunk_size=settings.notebook_chunk_size,
        chunk_overlap=settings.notebook_chunk_overlap,
        page_offsets=page_offsets,
    )
